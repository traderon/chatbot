from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
import openai
import os
from dotenv import load_dotenv
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from flask import session
from flask import send_file
from io import BytesIO
from datos import Estandar, Desempeño


load_dotenv()

app = Flask(__name__, static_folder='./build')
CORS(app)
#app.secret_key = os.environ.get("SECRET_KEY")
app.secret_key = "prueba123456789"
#client = openai(api_key=os.getenv("OPENAI_API_KEY"))

@app.route('/', defaults={'path': ''})
@app.route('/<path:path>')
def index(path):
    print("Entró a index")
    global preguntas
    global contexto
    global pregunta
    global grado
    global periodo
    global componente
    global estandar
    global desempeño
    global contador
    global messages
    session.clear()
    """GRADO 1:"""
    

    componente= ["Pensamiento numérico variacional", "Pensamiento Espacial Métrico", "Pensamiento Aleatorio"]
    barreras_familia=[
        "1. No hay acompañamiento en el desarrollo de los procesos educativos del estudiante por lo cual no se evidencian los resultados esperados",
        "2. No hay acompañamiento en el desarrollo de pautas de crianza, seguimiento, instrucciones y reconocimiento de figura de autoridad, por lo anterior afecta la convivencia escolar."
        ]
    barreras_docente=[
        "1. Creer que ciertos estudiantes no pueden lograr los mismos resultados que sus compañeros",
        "2. Negarse a modificar estrategias de enseñanza para adaptarse a las necesidades diversas del aula",
        "3. No considerar las dificultades individuales de los estudiantes y tratarlos con indiferencia",
        "4. Tener preferencias por ciertos alumnos y excluir o minimizar a otros",
        ]
    barreras_curriculares=[
        "1. Contenidos rígidos que no se adaptan a las necesidades de los estudiantes",
        "2. Falta de flexibilidad en los objetivos de aprendizaje para atender a estudiantes con dificultades específicas",
        "3. No hay ajustes curriculares en los DBA de acuerdo con las características del estudiante, su estilo y ritmo de aprendizaje"
        ]
    barreras_didacticas=[
        "1. No hay ajuste en Metodología, Espacio, Recurso y Comunicación",
        "2. No hay ajuste en Metodología y Recursos",
        "3. No hay ajustes en Espacio y Comunicación",
        "4. No hay ajustes en Metodología, Recursos, Comunicación y Tiempo",
        "5. No hay ajustes en Metodología, Recursos y Comunicación",
        "6. No hay ajustes en Metodología, Recursos y Tiempo",
        "7. No hay ajustes en Metodología, Espacio, Recursos",
        "8. No hay ajustes en Comunicación",
        "9. No hay ajustes en Metodología",
        "10. No hay ajustes en Tiempo",
        "11. No hay ajustes en Recursos",
        "12. No hay ajuste en Metodología, Espacio, Recurso, Tiempo y Comunicación"
        ]
    barreras_comunicativas=[
        "1. No hay ajuste en Comunicación Aumentativa",
        "2. No hay ajuste en Comunicación Alternativa",
        "3. No hay ajuste en Comunicación Aumentativa y Alternativa"
        ]
    barreras_fisicas=[
        "1. Mobiliario",
        "2. Ausencia de recursos tecnológicos",
        "3. Falta de apoyo, recursos y materiales",
        "4. Mobiliario, rampas, adecuaciones, etc",
        "5. Organización del espacio del aula y el plantel educativo",
        "6. Rampas",
        "7. Transporte o acceso educativo insuficiente"    
        ]
    session["grado"]=""
    session["barreras_familia"]=""
    session["barreras_docente"]=""
    session["barreras_curriculares"]=""
    session["barreras_didacticas"]=""
    session["barreras_comunicativas"]=""
    session["barreras_fisicas"]=""
    session["periodo"]=""
    session["componente"]=""
    session["estandar"]=""
    session["desempeño"]=""
    session["contador"]=0
    #pregunta=""
    #contexto={}
    #preguntas=['Cual/es el/los periodo/s a evaluar','Coloque el nivel obtenido por el estudiante: S(superior), A(Alto), B(Básico), Ba(Bajo)','Cuáles son las barreras actitudinales de la familia','Cuáles son las barreras actitudinales del docente','Cuáles son las barreras curriculares']
    session["preguntas"] = [
        "Seleccione el grado Grado 1 o Grado 2?",
        "Seleccione el período Primer período,Segundo período,Tercer período,Cuarto período?",
        "Seleccione el componente: \n\n"+componente[0]+" \n"+componente[1]+" \n"+componente[2],
        "Seleccione el estándar"+ session.get("grado", ""),
        "Seleccione el desempeño",
        "Selecciona cuales son las barreras actitudinales de la familia: \n\n"+barreras_familia[0]+" \n"+barreras_familia[1],
        "Selecciona cuales son las barreras actitudinales del docente: \n\n"+barreras_docente[0]+" \n"+barreras_docente[1]+" \n"+barreras_docente[2]+" \n"+barreras_docente[3],
        "Selecciona cuales son las barreras curriculares: \n\n"+barreras_curriculares[0]+" \n"+barreras_curriculares[1]+" \n"+barreras_curriculares[2],
        "Selecciona cuales son las barreras didacticas: \n\n"+barreras_didacticas[0]+" \n"+barreras_didacticas[1]+" \n"+barreras_didacticas[2]+" \n"+barreras_didacticas[3]+" \n"+barreras_didacticas[4]+" \n"+barreras_didacticas[5]+" \n"+barreras_didacticas[6]+" \n"+barreras_didacticas[7]+" \n"+barreras_didacticas[8]+" \n"+barreras_didacticas[9]+" \n"+barreras_didacticas[10]+" \n"+barreras_didacticas[11],
        "Selecciona cuales son las barreras comunicativas: \n\n"+barreras_comunicativas[0]+" \n"+barreras_comunicativas[1]+" \n"+barreras_comunicativas[2],
        "Selecciona cuales son las barreras físicas: \n\n"+barreras_fisicas[0]+" \n"+barreras_fisicas[1]+" \n"+barreras_fisicas[2]+" \n"+barreras_fisicas[3]+" \n"+barreras_fisicas[4]+" \n"+barreras_fisicas[5]+" \n"+barreras_fisicas[6]
        #"Coloque el nivel obtenido por el estudiante: S(superior), A(Alto), B(Básico), Ba(Bajo)",
        #"Cuáles son las barreras actitudinales de la familia",
        #"Cuáles son las barreras actitudinales del docente",
        #"Cuáles son las barreras curriculares"
    ]
    session["messages"]=[]
    session.modified = True
    
    print("SESSION:", dict(session))
    print("SET COOKIE?", app.session_interface.should_set_cookie(app, session))

    print("SESSION INDEX:", dict(session))
    session["contexto"] = {}
    session["pregunta"] = ""
    if path != "" and os.path.exists(app.static_folder + '/' + path):
        return send_from_directory(app.static_folder, path)
    else:
        return send_from_directory(app.static_folder, 'index.html')

def insertar_parrafo_despues(parrafo, texto):
    nuevo = OxmlElement("w:p")
    parrafo._p.addnext(nuevo)

    nuevo_parrafo = Paragraph(nuevo, parrafo._parent)
    nuevo_parrafo.add_run(texto)

    return nuevo_parrafo

@app.route("/chat", methods=["POST"])
def chat():
    openai.api_key = os.environ.get("OPENAI_SECRET_KEY")
    data = request.json
    message = data.get("message")
    #global contexto
    #global preguntas
    #global pregunta
    print("COOKIES:", request.cookies)
    print("SESSION:", dict(session))
    preguntas = session.get("preguntas", [])
    contexto = session.get("contexto", {})
    pregunta = session.get("pregunta", "")
    messages = session.get("messages", [])
    grado=session.get("grado", "")
    periodo=session.get("periodo", "")
    componente=session.get("componente", "")
    estandar=session.get("estandar", "")
    desempeño=session.get("desempeño", "")
    barreras_familia=session.get("barreras_familia", "")
    barreras_docente=session.get("barreras_docente", "")
    barreras_curriculares=session.get("barreras_curriculares", "")
    barreras_didacticas=session.get("barreras_didacticas", "")
    barreras_comunicativas=session.get("barreras_comunicativas", "")
    barreras_fisicas=session.get("barreras_fisicas", "")
    contador=session.get("contador", 0)
    print(preguntas)
    print(messages)
    messages.append({"role": "system", "content": "Eres un asistente educativo experto en ajustes razonables en matemáticas."})
    messages.append({
    "role": "user",
    "content": message
    })
    session.modified = True
    
    if len(preguntas)>=contador and any("Iniciar generación del ajuste razonable" in msg["content"] for msg in messages):
        #messages.append({"role": "user", "content": message})
        if len(preguntas)==contador:
            session["contador"] += 1
            contador=session.get("contador", 0)
            if "Selecciona cuales son las barreras físicas" in pregunta:
                session["barreras_fisicas"]=message
                barreras_fisicas=session.get("barreras_fisicas", "")
        elif len(preguntas)>contador:
            if pregunta !="" and "Seleccione el grado Grado 1 o Grado 2?" == pregunta:
                if message in "Grado 1" or message in "1" or message in "uno":
                    session["grado"]="Grado 1"
                elif message in "Grado 2" or message in "2" or message in "dos":
                    session["grado"]="Grado 2"
            elif pregunta !="" and "Seleccione el período Primer período,Segundo período,Tercer período,Cuarto período?" == pregunta:
                if message in "Primer período" or message in "1" or message in "Primer":
                    session["periodo"]="Primer período"
                elif message in "Segundo período" or message in "2" or message in "Segundo":
                    session["periodo"]="Segundo período"
            elif "Seleccione el componente" in pregunta:
                session["componente"]=message
            elif "Seleccione el estándar" in pregunta:
                session["estandar"]=message    
            elif "Selecciona cuales son las barreras actitudinales de la familia" in pregunta:
                session["barreras_familia"]=message
            elif "Selecciona cuales son las barreras actitudinales del docente" in pregunta:
                session["barreras_docente"]=message
            elif "Selecciona cuales son las barreras curriculares" in pregunta:   
                session["barreras_curriculares"]=message
            elif "Selecciona cuales son las barreras didacticas" in pregunta:
                session["barreras_didacticas"]=message
            elif "Selecciona cuales son las barreras comunicativas" in pregunta:
                session["barreras_comunicativas"]=message
            elif "Seleccione el desempeño" in pregunta:
                session["desempeño"]=message
            pregunta=preguntas[contador]
            if "Seleccione el estándar" in pregunta:
                estandares = Estandar[session["grado"]][session["periodo"]][session["componente"]]

                pregunta = "Seleccione el estándar:\n\n"
                
                for i, estandar in enumerate(estandares, start=1):
                    pregunta += f"{i}. {estandar}\n"
            elif "Seleccione el desempeño" in pregunta:
                if session["grado"]=='Grado 2':
                    desempeños = Desempeño[session["grado"]][session["periodo"]]["Pensamiento numérico variacional"]
                else:
                    desempeños = Desempeño[session["grado"]][session["periodo"]][session["componente"]]
                pregunta = "Seleccione el desempeño:\n\n"
                
                for i, (nivel,descripcion) in enumerate(desempeños.items(), start=1):
                    pregunta += f"{i}. {nivel}:{descripcion}\n"         
            # Guardar nuevamente en la sesión
            session["preguntas"] = preguntas
            session["contexto"] = contexto
            session["pregunta"] = pregunta
            session["messages"]= messages
            session["contador"] += 1
            print("SESSION:", dict(session))
            session.modified = True
            print("SESSION:", dict(session))
            #del preguntas[0]
            print(session["contexto"])
            return jsonify({
            "response": pregunta
            })    
    if contador>len(preguntas) and any("Iniciar generación del ajuste razonable" in msg["content"] for msg in messages):
        if len(preguntas)<3:
            contexto[pregunta]=message
        print(contexto)
        #documento = Document()
        doc = Document("PIAR.docx")
        
        for numeroCelda in range(1,4):
            celda = doc.tables[0].cell(1, numeroCelda)
            
            for i, p in enumerate(celda.paragraphs):
                print(i, p.text)
                if p.text.strip() == "PRIMER PERIODO" and numeroCelda==1:
                    insertar_parrafo_despues(
                        p,
                        periodo+"\n\n"+componente+"\n\n"+estandar
                    )
                    break
                if p.text.strip() == "Reformulación:" and numeroCelda==3:
                    insertar_parrafo_despues(
                        p,
                        desempeño
                    )
                    break
                if p.text.strip() == "FAMILIA:" and numeroCelda==2:
                    insertar_parrafo_despues(
                        p,
                        barreras_familia
                    )
                    
                if p.text.strip() == "DOCENTE:" and numeroCelda==2:
                    insertar_parrafo_despues(
                        p,
                        barreras_docente
                    )
                    
                if p.text.strip() == "BARRERAS CURRICULARES" and numeroCelda==2:
                    insertar_parrafo_despues(
                        p,
                        barreras_curriculares
                    )
                    
                if p.text.strip() == "BARRERA DIDÁCTICA" and numeroCelda==2:
                    insertar_parrafo_despues(
                        p,
                        barreras_didacticas
                    )
                    
                if p.text.strip() == "BARRERAS COMUNICATIVAS" and numeroCelda==2:
                    insertar_parrafo_despues(
                        p,
                        barreras_comunicativas
                    )   
                    
                if p.text.strip() == "BARRERAS FISICAS" and numeroCelda==2:
                    insertar_parrafo_despues(
                        p,
                        barreras_fisicas
                    ) 
                           
        doc.save("resultado.docx")
        # Crear tabla de 2 columnas
        #tabla = documento.add_table(rows=1, cols=3)
        #tabla.style = 'Table Grid'
        
        # Encabezados
        #encabezado = tabla.rows[0].cells
        #encabezado[0].text = contexto["Cual/es el/los periodo/s a evaluar"]
        #encabezado[1].text = 'Barreras actitudinales del aprendizaje'
        #encabezado[2].text = 'PRIMER PERIODO Y SEGUNDO PERFIODO'

        
        # Agregar datos
        
        #fila = tabla.add_row().cells
        #fila[0].text = contexto["Coloque el nivel obtenido por el estudiante: S(superior), A(Alto), B(Básico), Ba(Bajo)"]
        #fila[1].text = 'Familia: \n'+str(list(contexto.values())[2])+'\n Docente: \n'+str(list(contexto.values())[3])+'\n Barreras curriculares: \n'+str(list(contexto.values())[4])
        #fila[2].text = 'ADAPTACIÓN CURRICULAR: \n'+"""Priorización:  .por dos periodos académicos 
        #Desarrollar el reconocimiento visual y manipulativo de los números del 1 al 10 mediante el uso de material concreto, pictogramas y rutinas repetitivas, permitiendo que la estudiante identifique cada número y lo asocie con una cantidad correspondiente, sin requerir lenguaje oral o simbólico complejo.Reformulación: 
        #Reconoce y compara cantidades simples (mitad, entero, vacío, lleno) usando material concreto y apoyos visuales, sin necesidad de lenguaje estructurado. Identifica cuándo una cantidad es “una parte” o “todas las partes” en diferentes situaciones manipulativas.
        #"""
        
        #documento.save('reporte.docx')
        #return jsonify({
        #"response": "Gracias por tus respuestas"
        # })
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name="reporte.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
         
    """
    response = openai.ChatCompletion.create(
            model="gpt-4",  # Modelo que deseas usar
            #model="gpt-3.5-turbo",
            messages=messages,
            temperature=0.7, 
            max_tokens=2000  # Límite de tokens para la respuesta
        )
        #print(contexto)
        # Obtener la respuesta generada
    bot_reply = response['choices'][0]['message']['content']
   


    return jsonify({
        "response": bot_reply
    })
    """
if __name__ == "__main__":
    app.run(debug=True)